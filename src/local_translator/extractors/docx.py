from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from local_translator.extractors.base import BaseExtractor, ExtractedDocument


@dataclass(slots=True)
class ParagraphSnapshot:
    segment_indexes: list[int]


@dataclass(slots=True)
class TableSnapshot:
    rows: list[list[list[ParagraphSnapshot]]]


def _iter_blocks(doc: DocumentObject):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _extract_paragraph(paragraph: Paragraph, segments: list[str]) -> ParagraphSnapshot:
    segment_indexes: list[int] = []
    if paragraph.runs:
        for run in paragraph.runs:
            segments.append(run.text)
            segment_indexes.append(len(segments) - 1)
    else:
        segments.append(paragraph.text)
        segment_indexes.append(len(segments) - 1)
    return ParagraphSnapshot(segment_indexes=segment_indexes)


def _translated_value(extracted: ExtractedDocument, translated_segments: list[str], index: int) -> str:
    if index < len(translated_segments):
        return translated_segments[index]
    if index < len(extracted.segments):
        return extracted.segments[index]
    return ""


def _apply_paragraph_translation(
    paragraph: Paragraph,
    snapshot: ParagraphSnapshot,
    extracted: ExtractedDocument,
    translated_segments: list[str],
) -> None:
    values = [
        _translated_value(extracted=extracted, translated_segments=translated_segments, index=index)
        for index in snapshot.segment_indexes
    ]

    if paragraph.runs and len(paragraph.runs) == len(values):
        for run, value in zip(paragraph.runs, values):
            run.text = value
        return

    paragraph.text = "".join(values)


class DocxExtractor(BaseExtractor):
    suffixes = (".docx",)

    def extract(self, file_path: Path) -> ExtractedDocument:
        doc = Document(str(file_path))
        segments: list[str] = []
        blocks: list[dict[str, Any]] = []

        for block in _iter_blocks(doc):
            if isinstance(block, Paragraph):
                snapshot = _extract_paragraph(block, segments)
                blocks.append(
                    {
                        "type": "paragraph",
                        "style": block.style.name if block.style is not None else None,
                        "paragraph": snapshot,
                    }
                )
                continue

            table_rows: list[list[list[ParagraphSnapshot]]] = []
            for row in block.rows:
                row_cells: list[list[ParagraphSnapshot]] = []
                for cell in row.cells:
                    cell_paragraphs = [_extract_paragraph(paragraph, segments) for paragraph in cell.paragraphs]
                    row_cells.append(cell_paragraphs)
                table_rows.append(row_cells)

            blocks.append({"type": "table", "table": TableSnapshot(rows=table_rows)})

        return ExtractedDocument(file_path=file_path, segments=segments, metadata={"blocks": blocks})

    def reconstruct(self, extracted: ExtractedDocument, translated_segments: list[str], output_path: Path) -> None:
        doc = Document(str(extracted.file_path))
        blocks_meta = extracted.metadata.get("blocks", []) if extracted.metadata else []

        for block, block_meta in zip(_iter_blocks(doc), blocks_meta):
            if isinstance(block, Paragraph) and block_meta.get("type") == "paragraph":
                _apply_paragraph_translation(
                    paragraph=block,
                    snapshot=block_meta["paragraph"],
                    extracted=extracted,
                    translated_segments=translated_segments,
                )
                continue

            if isinstance(block, Table) and block_meta.get("type") == "table":
                table_meta: TableSnapshot = block_meta["table"]
                for row, row_meta in zip(block.rows, table_meta.rows):
                    for cell, cell_meta in zip(row.cells, row_meta):
                        for paragraph, paragraph_meta in zip(cell.paragraphs, cell_meta):
                            _apply_paragraph_translation(
                                paragraph=paragraph,
                                snapshot=paragraph_meta,
                                extracted=extracted,
                                translated_segments=translated_segments,
                            )

        doc.save(str(output_path))
